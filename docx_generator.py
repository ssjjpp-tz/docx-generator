"""
Word文档生成器
支持论文翻译、格式化输出
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os


class DocxGenerator:
    """Word文档生成器"""
    
    def __init__(self):
        self.doc = Document()
        # 设置默认字体
        self._set_default_font()
    
    def _set_default_font(self):
        """设置文档默认字体"""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def add_title(self, text, level=1):
        """添加标题"""
        heading = self.doc.add_heading(text, level=level)
        return heading
    
    def add_paragraph(self, text, style=None):
        """添加段落"""
        p = self.doc.add_paragraph(text, style=style)
        return p
    
    def add_translation(self, original, translated):
        """添加翻译对照（原文+译文）"""
        # 原文
        p1 = self.doc.add_paragraph()
        p1.add_run("原文：").bold = True
        p1.add_run(original)
        
        # 译文
        p2 = self.doc.add_paragraph()
        p2.add_run("译文：").bold = True
        p2.add_run(translated)
        
        # 空行分隔
        self.doc.add_paragraph()
    
    def add_paper_section(self, title, content):
        """添加论文章节"""
        # 章节标题
        self.add_title(title, level=2)
        # 内容
        self.add_paragraph(content)
        # 空行
        self.doc.add_paragraph()
    
    def add_page_break(self):
        """添加分页"""
        self.doc.add_page_break()
    
    def save(self, filepath):
        """保存文档"""
        # 确保目录存在
        os.makedirs(os.path.dirname(os.path.abspath(filepath)) if os.path.dirname(filepath) else '.', exist_ok=True)
        self.doc.save(filepath)
        return filepath
    
    @classmethod
    def generate_from_text(cls, text, output_path, title=None):
        """从纯文本生成文档"""
        gen = cls()
        
        if title:
            gen.add_title(title)
        
        # 按段落分割
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            if para.strip():
                gen.add_paragraph(para.strip())
        
        return gen.save(output_path)


if __name__ == "__main__":
    # 测试
    gen = DocxGenerator()
    gen.add_title("测试文档", level=1)
    gen.add_paragraph("这是一个测试段落。")
    gen.save("/tmp/test_output.docx")
    print("测试文档已生成: /tmp/test_output.docx")
